#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Diego Minguzzi 2025
from .document_reader import DocumentReader
from .word_document_profile import WordDocumentProfile
from entities.doc import Document, Figure, Metadata, Paragraph, Section

import docx
import logging as log
import io
import os
import PIL
import typing

class WordDocumentReader(DocumentReader):

  def __init__( self, stripParagraphs: bool= True):
    super().__init__()
    self._docProfile= WordDocumentProfile()
    self._parStyleRank= self._docProfile.paragraphStyleRank

    self._stripParagraphs= stripParagraphs

  def readDocument(self, filepath: str) -> Document:
    metadata = Metadata.createFromFilepath( filepath)

    tmpTitle = os.path.splitext(os.path.basename(filepath))[0]

    indxParagraph=0

    wordDoc = docx.Document(filepath)
    paragraphs= wordDoc.paragraphs

    if len(paragraphs)==0:
      return Document( '', metadata)

    rootSection= Section()

    par= paragraphs[indxParagraph]
    parStyleName= par.style.name

    while indxParagraph < len(paragraphs):
      indx = self._traverseSection( wordDoc=wordDoc,
                                    paragraphs=paragraphs,
                                    indxParagraph=indxParagraph,
                                    parentSection=rootSection)
      indxParagraph= indx

    return Document( tmpTitle, metadata, rootSection)

  #------------------------------------------------------------------------------------------------
  def _traverseSection( self,
                        wordDoc: docx.Document,
                        paragraphs,
                        indxParagraph: int,
                        parentSection: Section,
                        parentRank:int=0) -> int:
    """ Traverses a section.
        Tries to add all sections as subsections to parent, that have rank higher than parentRank.
        Returns the index of the next paragraph to be processed.
    """
    indx = indxParagraph

    par= paragraphs[indx]
    parStyleName= par.style.name

    if self._parStyleRank.isTextRank( parStyleName):
      (indx, paragraph) = self._traverseParagraph( wordDoc, paragraphs, indx)
      currSection= parentSection.createSubsection( paragraph=paragraph)

    else:
      while indx<len(paragraphs):
        par= paragraphs[indx]
        parStyleName= par.style.name
        assert( not self._parStyleRank.isTextRank( parStyleName))

        sectRank = self._parStyleRank.getRank( parStyleName)
        if sectRank <= parentRank:
          return indx

        sectionTitle= par.text

        if (indx+1) >= len(paragraphs):
          currSection= parentSection.createSubsection( title= sectionTitle, paragraph=Paragraph())
          return indx+1 # TODO
        else:
          nextIndx= indx+1
          nextPar= paragraphs[ nextIndx]
          nextParName= nextPar.style.name
          nextRank = self._parStyleRank.getRank( nextParName)

          if self._parStyleRank.isTextRank( nextParName):
            (indx, paragraph) = self._traverseParagraph( wordDoc, paragraphs, nextIndx)
            if indx>=len(paragraphs):
              currSection= parentSection.createSubsection( title= sectionTitle, paragraph=paragraph)
              return indx

            nextIndx= indx
            nextPar= paragraphs[ nextIndx]
            nextParName= nextPar.style.name
            assert( not self._parStyleRank.isTextRank( nextParName))
            nextRank = self._parStyleRank.getRank( nextParName)
          else:
            paragraph= Paragraph()

          currSection= parentSection.createSubsection( title= sectionTitle, paragraph=paragraph)

          nextRank = self._parStyleRank.getRank( nextParName)
          if sectRank < nextRank:
            indx = self._traverseSection( wordDoc, paragraphs, nextIndx, currSection, sectRank)
          elif nextRank <= parentRank:
            return nextIndx
          else:
            indx = nextIndx

    return indx

  #------------------------------------------------------------------------------------------------
  def _traverseParagraph(self,
                         wordDoc: docx.Document,
                         paragraphs,
                         indxParagraph: int) -> typing.Tuple[int, Paragraph]:
    """ Traverses a paragraph.
        Stops when it finds a header item.
        Returns the index of non paragraph item and the paragraph object."""
    indx= indxParagraph
    textRows= list()
    allFigures= list()

    hasFoundNonEmptyRow= False

    while indx < len(paragraphs):
      parStyleName= paragraphs[indx].style.name
      paragraph = paragraphs[indx]
      if self._parStyleRank.isTextRank( parStyleName):
        rowText= paragraph.text

        if self._stripParagraphs:
          isRowEmpty= (0==len(rowText.strip()))
          if (not hasFoundNonEmptyRow) and isRowEmpty:
            indx += 1
            continue
          if not isRowEmpty:
            hasFoundNonEmptyRow= True

        textRows.append(rowText)

        figures= self._retrieveFigures( wordDoc, paragraph)
        if len(figures)>0:
          allFigures.extend( figures)

        indx += 1
      else:
        break

    if self._stripParagraphs:
      while len(textRows)>0 and not textRows[-1].strip():
        textRows.pop()

    log.debug(f'Retrieved num figures: {len(allFigures)}')
    return (indx, Paragraph( text=textRows, figures=allFigures))

  #--------------------------------------------------------------------------------------------------
  def _retrieveFigures(self,
                      wordDoc,
                      paragraph) -> typing.List[Figure]:
    figures= list()

    for run in paragraph.runs:
      # Check if the run contains an inline shape (potential picture)
      if run._element.xpath('.//w:drawing'):
          log.debug('Figure detected in this paragraph.')

          for drawing in run._element.xpath('.//w:drawing'):

            # Extract the image data from the drawing element
            blip_elements = drawing.xpath('.//a:blip')
            for blip in blip_elements:
              embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
              if embed:
                # Access the image part from the document
                image_part = wordDoc.part.related_parts[embed]
                image_data = image_part.blob

                image = PIL.Image.open( io.BytesIO(image_data))
                figures.append( Figure( image=image))

    return figures
