import argparse
import docx
from PIL import Image
from io import BytesIO
"""
python ${REPO_ROOT}/poc_read_word_docx.py --word-file ${TAI_DATASET_ROOT}/cobot_kb/PARobot-plugin-3.0C1.docx
python ${REPO_ROOT}/poc_read_word_docx.py --word-file ${TAI_DATASET_ROOT}/cobot_kb/Guidance_Programming_Language.docx
python ${REPO_ROOT}/src_py/poc_doc_reader/poc_read_word_docx.py --word-file ${TAI_DATASET_ROOT}/cobot_kb/Example.docx
"""

class CliArgs:
    def __init__(self):
        parser = argparse.ArgumentParser(description="Process a word file.")
        parser.add_argument("--word-file",type=str,required=True,help="Path to the word file")
        args = parser.parse_args()
        self._wordFile = args.word_file

    @property
    def wordFile(self):
        return self._wordFile

def traverseDocument(filepath):
    # Open the Word document
    doc = docx.Document(filepath)

    # Iterate through all paragraphs
    for paragraph in doc.paragraphs:
        # Check for specific heading styles
        if paragraph.style.name == 'Heading 1':
            print(f"Header 1: {paragraph.text}")
        elif paragraph.style.name == 'Heading 2':
            print(f"Header 2: {paragraph.text}")
        elif paragraph.style.name == 'Heading 3':
            print(f"Header 3: {paragraph.text}")
        elif paragraph.style.name == 'Heading 4':
            print(f"Header 4: {paragraph.text}")
        elif paragraph.style.name == 'Heading 5':
            print(f"Header 5: {paragraph.text}")
        else:
            # Print normal text if it's not a heading
            if 0!=len(str(paragraph.text).strip()):
              print(f"Unknown style: {paragraph.style.name}: content:{paragraph.text}")

        # Detect pictures in the paragraph
        for run in paragraph.runs:
            # Check if the run contains an inline shape (potential picture)
            if run._element.xpath('.//w:drawing'):
                print("Picture detected in this paragraph!")
                for drawing in run._element.xpath('.//w:drawing'):
                  print(f'drawing type: {type(drawing)}')
                  # docx.oxml.drawing.CT_Drawing
                  # Extract the image data from the drawing element
                  blip_elements = drawing.xpath('.//a:blip')
                  for blip in blip_elements:
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                      # Access the image part from the document
                      image_part = doc.part.related_parts[embed]
                      image_data = image_part.blob

                      # Save the image to a file
                      image = Image.open(BytesIO(image_data))
                      image_filename = f"extracted_image_{id(image)}.png"
                      # image.save(image_filename)
                      print(f"Image filename {image_filename}")


def main():
    pass  # Replace with your code
    cliArgs = CliArgs()
    traverseDocument( cliArgs.wordFile)

if __name__ == "__main__":
    main()
